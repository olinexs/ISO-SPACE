from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re
import os

def detect_placeholders(doc):
    """
    Detects all placeholders enclosed in curly braces {} within the document,
    preserving top-to-bottom order.
    """
    placeholders = []

    # Search in paragraphs (top-to-bottom)
    for paragraph in doc.paragraphs:
        matches = re.findall(r"\{.*?\}", paragraph.text)
        placeholders.extend(matches)  # Add matches to the list

    # Search in tables row by row (top-to-bottom)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    matches = re.findall(r"\{.*?\}", paragraph.text)
                    placeholders.extend(matches)  # Add matches to the list

    # Return unique placeholders in the order they were found
    return list(dict.fromkeys(placeholders))  # Removes duplicates while preserving order



def replace_placeholder_with_image(cell, placeholder, image_path, width_in_inches=1):
    """
    Replaces a placeholder with an image in a cell, preserving surrounding text.
    """
    if placeholder in cell.text:
        for paragraph in cell.paragraphs:
            if placeholder in paragraph.text:
                parts = paragraph.text.split(placeholder)
                paragraph.text = parts[0]  # Text before the placeholder
                run = paragraph.add_run()
                run.add_picture(image_path, width=Inches(width_in_inches))
                if len(parts) > 1:
                    paragraph.add_run(parts[1])  # Text after the placeholder
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


def replace_placeholder_with_text(cell, placeholder, value):
    """
    Replaces a placeholder with text in a cell, preserving surrounding text.
    """
    if placeholder in cell.text:
        for paragraph in cell.paragraphs:
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, value)


from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def replace_placeholder_in_header_footer(header_footer, placeholder, image_path, width_in_inches):
    for paragraph in header_footer.paragraphs:
        if placeholder in paragraph.text:
            parts = paragraph.text.split(placeholder)
            paragraph.text = parts[0]
            run = paragraph.add_run()
            run.add_picture(image_path, width=Inches(width_in_inches))
            if len(parts) > 1:
                paragraph.add_run(parts[1])
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

def generate_document_from_template(
    template_path, output_path, replacements, image_replacements
):
    doc = Document(template_path)

    # Define the path for the logo image
    logo_image_path = "logo/path.png" 
    logo_placeholder = "{logo}"

    # Replace placeholders in headers and footers
    for section in doc.sections:
        if section.header:
            replace_placeholder_in_header_footer(section.header, logo_placeholder, logo_image_path, 1.5)

        if section.footer:
            for placeholder, image_path in image_replacements.items():
                replace_placeholder_in_header_footer(section.footer, placeholder, image_path, 1.5)

    # Replace placeholders in body tables and paragraphs
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for placeholder, value in replacements.items():
                    replace_placeholder_with_text(cell, placeholder, value)
                for placeholder, image_path in image_replacements.items():
                    replace_placeholder_with_image(cell, placeholder, image_path)

    for paragraph in doc.paragraphs:
        for placeholder, value in replacements.items():
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, value)
        for placeholder, image_path in image_replacements.items():
            if placeholder in paragraph.text:
                parts = paragraph.text.split(placeholder)
                paragraph.text = parts[0]
                run = paragraph.add_run()
                run.add_picture(image_path, width=Inches(1.5))
                if len(parts) > 1:
                    paragraph.add_run(parts[1])

    doc.save(output_path)
    print(f"Document saved to: {output_path}")


def main():
    # Prompt for template path
    template_path = input("Enter the path to the template document (.docx): ").strip()

    # Load the document to detect placeholders
    doc = Document(template_path)
    placeholders = detect_placeholders(doc)
    print("\nDetected placeholders:")
    for placeholder in placeholders:
        print(f"  - {placeholder}")

    # Collect input for replacements
    replacements = {}
    image_replacements = {}
    print("\nEnter values for each placeholder:")

    for placeholder in placeholders:
        if "sign" in placeholder.lower():
            # Assume these placeholders are for images
            image_path = input(f"path to {placeholder.strip('{}')}: ").strip()
            image_replacements[placeholder] = image_path
        else:
            # Assume these placeholders are for text
            text_value = input(f"{placeholder.strip('{}')}: ").strip()
            replacements[placeholder] = text_value

    # Output file path
    output_path = input("\nEnter the path to save the generated document: ").strip()

    # Generate the document
    generate_document_from_template(template_path, output_path, replacements, image_replacements)


if __name__ == "__main__":
    main()
