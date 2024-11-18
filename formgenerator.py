from docx import Document
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os

# Define the WordprocessingML namespace
NAMESPACE = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}


def apply_borders_to_row(row):
    """Apply default borders to all cells in a row."""
    for cell in row.cells:
        tcPr = cell._element.get_or_add_tcPr()
        borders = OxmlElement('w:tcBorders')

        for side in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{side}')
            border.set(qn('w:val'), 'single')  
            border.set(qn('w:sz'), '4')  
            border.set(qn('w:space'), '0') 
            border.set(qn('w:color'), 'auto')  
            borders.append(border)

        existing_borders = tcPr.find('w:tcBorders', NAMESPACE)
        if existing_borders is not None:
            tcPr.remove(existing_borders)
        tcPr.append(borders)


def get_table_headers(table):
    """Retrieve headers from the first row of the table."""
    return [cell.text.strip() for cell in table.rows[0].cells]


def collect_data_from_user(headers):
    """Prompt the user to input data for each row dynamically based on headers."""
    print("\nInput table data:")
    table_data = []
    while True:
        row_data = []
        print("\nEnter data for a new row (or type 'done' to finish):")
        for header in headers:
            value = input(f"  {header}: ")
            if value.lower() == 'done':
                return table_data
            row_data.append(value)
        table_data.append(row_data)
        
def replace_placeholder_with_image(paragraphs, placeholder, image_path, width_inches=1):
    """Replace a placeholder in paragraphs with an image."""
    for paragraph in paragraphs:
        if placeholder in paragraph.text:
            paragraph.clear()
            run = paragraph.add_run()
            run.add_picture(image_path, width=Inches(width_inches))


def generate_document(template_path, logo_path, company_name, signature_paths, output_path):
    output_dir = os.path.dirname(output_path)
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    doc = Document(template_path)

    for section in doc.sections:
        header = section.header
        for paragraph in header.paragraphs:
            if "{logo}" in paragraph.text:
                paragraph.clear()
                run = paragraph.add_run()
                run.add_picture(logo_path, width=Inches(1.5))

            if "{Company Name}" in paragraph.text:
                paragraph.text = paragraph.text.replace("{Company Name}", company_name)
                
    for paragraph in doc.paragraphs:
        replace_placeholder_with_image(paragraphs=[paragraph], placeholder="{issued_sign}", image_path=signature_paths.get("issued_sign"))
        replace_placeholder_with_image(paragraphs=[paragraph], placeholder="{approved_sign}", image_path=signature_paths.get("approved_sign"))


    if len(doc.tables) > 0:
        table = doc.tables[0]
        headers = get_table_headers(table)
        print(f"Detected table headers: {headers}")
        table_data = collect_data_from_user(headers)

        for i, row_data in enumerate(table_data):
            if i + 1 < len(table.rows):
                row = table.rows[i + 1]
            else:
                row = table.add_row()

            apply_borders_to_row(row)

            for j, cell_data in enumerate(row_data):
                if j < len(row.cells):  
                    row.cells[j].text = cell_data
    else:
        raise ValueError("No tables found in the document template.")

    doc.save(output_path)
    print(f"Document saved to: {output_path}")


# Example usage
template_path = ""
logo_path = ""
output_path = ""
company_name = ""

signature_paths = {
    "issued_sign": "",
    "approved_sign": ""
}

generate_document(template_path, logo_path, company_name, signature_paths, output_path)
